#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Pandoc Word文档处理工具 - 增强版 (支持加点字检测，图片处理已禁用)

使用pandoc将Word文档转换为模型可读的纯文本内容，支持：
1. 文档文本转换 (Pandoc)
2. 大模型API调用 (文档结构解析)
3. 着重号检测 (加点字标记保留)
4. 连续短横线转中文破折号

注：图片提取和内容分析功能已禁用以提高运行效率
注：加点字转换为HTML功能已移除，标记保留供后续模型处理

依赖安装：
1. 确保系统已安装pandoc: https://pandoc.org/installing.html
2. 安装python-docx: pip install python-docx (仅用于加点字预处理)
3. 安装其他依赖: pip install requests

使用方法：
1. 运行脚本处理Word文档
2. 检测并保留加点字格式标记
3. 生成最终的解析结果
"""

import subprocess
import requests
import json
import time
import os
import re
import zipfile
import tempfile
from datetime import datetime
from pathlib import Path
from PIL import Image
import base64
from io import BytesIO
from collections import defaultdict

# 特殊格式识别依赖
try:
    from docx import Document
    from docx.enum.text import WD_UNDERLINE
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
    print("✅ python-docx库可用，支持特殊格式识别")
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx库不可用，将跳过特殊格式识别功能")
    print("   安装命令: pip install python-docx")


class PandocWordProcessor:
    def __init__(self):
        self.api_key = "baf9ea42-7e17-4df6-9a22-90127ac8220e"
        self.base_url = "https://ark.cn-beijing.volces.com/api"

        def _check_pandoc(self):
            """检查pandoc是否可用"""
            try:
                result = subprocess.run(['pandoc', '--version'],
                                        capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    print(f"✅ Pandoc可用: {result.stdout.split()[1]}")
                    return True
                else:
                    print(f"❌ Pandoc检查失败: {result.stderr}")
                    return False
            except FileNotFoundError:
                print("❌ 未找到pandoc命令")
                return False
            except subprocess.TimeoutExpired:
                print("❌ Pandoc检查超时")
                return False
            except Exception as e:
                print(f"❌ Pandoc检查异常: {e}")
                return False

        # 检查pandoc是否可用
        self.pandoc_available = _check_pandoc(self)
        if not self.pandoc_available:
            print("⚠️ 警告: pandoc未安装或不在PATH中")
            print("请访问 https://pandoc.org/installing.html 安装pandoc")

        def _init_format_styles(self):
            """初始化格式样式映射"""
            if not DOCX_AVAILABLE:
                return

            # 下划线样式映射 - 安全地添加下划线样式
            self.underline_styles = {}
            styles_to_add = [
                (getattr(WD_UNDERLINE, 'SINGLE', None), "单下划线"),
                (getattr(WD_UNDERLINE, 'DOUBLE', None), "双下划线"),
                (getattr(WD_UNDERLINE, 'THICK', None), "粗下划线"),
                (getattr(WD_UNDERLINE, 'DOTTED', None), "点状下划线"),
                (getattr(WD_UNDERLINE, 'DASH', None), "虚线下划线"),
                (getattr(WD_UNDERLINE, 'DOT_DASH', None), "点划线下划线"),
                (getattr(WD_UNDERLINE, 'DOT_DOT_DASH', None), "点点划线下划线"),
                (getattr(WD_UNDERLINE, 'WAVY', None), "波浪线下划线"),
                (getattr(WD_UNDERLINE, 'DOTTED_HEAVY', None), "粗点状下划线"),
                (getattr(WD_UNDERLINE, 'DASH_HEAVY', None), "粗虚线下划线"),
                (getattr(WD_UNDERLINE, 'WAVY_HEAVY', None), "粗波浪线下划线"),
                (getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None), "双波浪线下划线")
            ]

            for style_enum, style_name in styles_to_add:
                if style_enum is not None:
                    self.underline_styles[style_enum] = style_name

            print(f"📋 初始化了 {len(self.underline_styles)} 种下划线样式识别")

        # 特殊格式识别功能初始化
        self.format_detection_enabled = DOCX_AVAILABLE
        self.special_formatted_text = []
        self.format_statistics = defaultdict(int)
        self.paragraph_formatting = []  # 新增：存储段落格式信息

        if self.format_detection_enabled:
            _init_format_styles(self)

    def _analyze_text_formatting(self, run, para_index=0, run_index=0):
        """工具函数：分析文本片段的格式"""
        if not DOCX_AVAILABLE:
            return []

        formats = []
        font = run.font

        def _check_emphasis_mark(self, run):
            """工具函数：检查着重号（加点字）"""
            try:
                run_xml = run._element
                em_elements = run_xml.xpath('.//w:em', namespaces={
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

                if em_elements:
                    em_val = em_elements[0].get(qn('w:val'))
                    emphasis_types = {
                        'dot': '点',
                        'comma': '逗号',
                        'circle': '圆圈',
                        'underDot': '下点'
                    }
                    return emphasis_types.get(em_val, f'着重号({em_val})')
            except:
                pass

            return None

        # 下划线检查
        if font.underline:
            underline_style = self.underline_styles.get(font.underline, f"未知下划线样式({font.underline})")
            formats.append(f"下划线: {underline_style}")

            # 特别标记波浪线和点状线
            wavy_styles = [style for style in [
                getattr(WD_UNDERLINE, 'WAVY', None),
                getattr(WD_UNDERLINE, 'WAVY_HEAVY', None),
                getattr(WD_UNDERLINE, 'WAVY_DOUBLE', None)
            ] if style is not None]

            dotted_styles = [style for style in [
                getattr(WD_UNDERLINE, 'DOTTED', None),
                getattr(WD_UNDERLINE, 'DOTTED_HEAVY', None)
            ] if style is not None]

            if font.underline in wavy_styles:
                formats.append("⚠️ 波浪线格式")
            elif font.underline in dotted_styles:
                formats.append("⚠️ 点状线格式")

        # 上标下标
        if font.superscript:
            formats.append("上标")
        if font.subscript:
            formats.append("下标")

        # 删除线
        if font.strike:
            formats.append("删除线")

        # 粗体斜体
        if font.bold:
            formats.append("粗体")
        if font.italic:
            formats.append("斜体")

        # 字体颜色
        if font.color and font.color.rgb:
            try:
                rgb_val = font.color.rgb
                if hasattr(rgb_val, '__int__'):
                    color_hex = f"#{int(rgb_val):06x}"
                else:
                    color_hex = str(rgb_val)
                formats.append(f"字体颜色: {color_hex}")
            except Exception:
                formats.append("字体颜色: 特殊颜色")

        # 字体大小
        if font.size:
            try:
                size_pt = font.size.pt
                formats.append(f"字号: {size_pt}磅")
            except:
                formats.append("字号: 自定义大小")

        # 字体名称
        if font.name:
            formats.append(f"字体: {font.name}")

        # 检查着重号
        emphasis_mark = _check_emphasis_mark(self, run)
        if emphasis_mark:
            formats.append(f"着重号: {emphasis_mark}")

        # 统计格式使用
        for fmt in formats:
            self.format_statistics[fmt] += 1

        # 保存特殊格式的文本
        if formats and run.text.strip():
            self.special_formatted_text.append({
                'text': run.text,
                'paragraph': para_index,
                'run': run_index,
                'formats': formats
            })

        return formats

    def _analyze_paragraph_formatting(self, paragraph, para_index=0):
        """分析段落的格式，包括首行缩进"""
        if not DOCX_AVAILABLE:
            return []

        para_formats = []

        try:
            # 获取段落格式
            para_format = paragraph.paragraph_format

            # 检查首行缩进
            if para_format.first_line_indent:
                indent_value = para_format.first_line_indent
                # 转换为磅数（如果可能）
                try:
                    indent_pt = indent_value.pt if hasattr(indent_value, 'pt') else None
                    if indent_pt and indent_pt > 0:
                        para_formats.append(f"首行缩进: {indent_pt:.1f}磅")
                        # 标记这是一个重要的格式信息
                        para_formats.append("⚠️ 首行缩进段落")
                except:
                    para_formats.append("首行缩进: 自定义值")
                    para_formats.append("⚠️ 首行缩进段落")

            # 检查左缩进
            if para_format.left_indent:
                try:
                    left_indent_pt = para_format.left_indent.pt if hasattr(para_format.left_indent, 'pt') else None
                    if left_indent_pt and left_indent_pt > 0:
                        para_formats.append(f"左缩进: {left_indent_pt:.1f}磅")
                except:
                    para_formats.append("左缩进: 自定义值")

            # 检查右缩进
            if para_format.right_indent:
                try:
                    right_indent_pt = para_format.right_indent.pt if hasattr(para_format.right_indent, 'pt') else None
                    if right_indent_pt and right_indent_pt > 0:
                        para_formats.append(f"右缩进: {right_indent_pt:.1f}磅")
                except:
                    para_formats.append("右缩进: 自定义值")

            # 检查对齐方式
            if para_format.alignment:
                alignment_names = {
                    0: "左对齐",
                    1: "居中",
                    2: "右对齐",
                    3: "两端对齐",
                    4: "分散对齐"
                }
                alignment_name = alignment_names.get(para_format.alignment, f"对齐方式{para_format.alignment}")
                para_formats.append(f"对齐: {alignment_name}")

            # 统计格式使用
            for fmt in para_formats:
                self.format_statistics[fmt] += 1

        except Exception as e:
            print(f"  ⚠️ 段落格式分析失败: {e}")

        return para_formats

    def extract_format_analysis(self, docx_path):
        """提取文档的格式分析信息"""
        if not self.format_detection_enabled:
            print("🚫 格式检测功能未启用，跳过格式分析")
            return None

        print("🔍 开始分析文档格式...")

        try:
            doc = Document(docx_path)
            paragraph_count = 0

            # 重置统计信息
            self.special_formatted_text = []
            self.format_statistics = defaultdict(int)
            self.paragraph_formatting = []  # 重置段落格式信息

            for para in doc.paragraphs:
                paragraph_count += 1

                # 分析段落格式（首行缩进等）
                para_formats = self._analyze_paragraph_formatting(para, paragraph_count)
                if para_formats:
                    # 保存段落格式信息，包括段落文本
                    para_text = para.text.strip()
                    if para_text:  # 只保存非空段落
                        self.paragraph_formatting.append({
                            'paragraph_index': paragraph_count,
                            'text': para_text,
                            'formats': para_formats,
                            'has_first_line_indent': any('首行缩进' in fmt for fmt in para_formats),
                            'is_centered': any('对齐: 居中' in fmt for fmt in para_formats),
                            'is_right_aligned': any('对齐: 右对齐' in fmt for fmt in para_formats)
                        })

                # 分析段落中的文本格式
                for run_index, run in enumerate(para.runs):
                    self._analyze_text_formatting(run, paragraph_count, run_index)

            # 分析表格
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        for para in cell.paragraphs:
                            for run_index, run in enumerate(para.runs):
                                location = f"表格{table_count}-行{row_index}-列{cell_index}"
                                self._analyze_text_formatting(run, location, run_index)

            # 统计首行缩进段落、居中段落和居右段落
            indent_paragraphs = len([p for p in self.paragraph_formatting if p['has_first_line_indent']])
            centered_paragraphs = len([p for p in self.paragraph_formatting if p['is_centered']])
            right_aligned_paragraphs = len([p for p in self.paragraph_formatting if p['is_right_aligned']])

            print(f"✅ 格式分析完成: {paragraph_count}个段落, {table_count}个表格")
            print(f"📊 发现 {len(self.special_formatted_text)} 个包含特殊格式的文本片段")
            print(f"📏 发现 {indent_paragraphs} 个包含首行缩进的段落")
            print(f"📐 发现 {centered_paragraphs} 个居中对齐的段落")
            print(f"📑 发现 {right_aligned_paragraphs} 个居右对齐的段落")

            return {
                'total_paragraphs': paragraph_count,
                'total_tables': table_count,
                'special_format_count': len(self.special_formatted_text),
                'paragraph_format_count': len(self.paragraph_formatting),
                'indent_paragraph_count': indent_paragraphs,
                'centered_paragraph_count': centered_paragraphs,
                'right_aligned_paragraph_count': right_aligned_paragraphs,
                'format_statistics': dict(self.format_statistics)
            }

        except Exception as e:
            print(f"❌ 格式分析失败: {e}")
            return None

    def preprocess(self, docx_path):
        print("🔍 预处理...")

        try:
            # 导入预处理器
            import zipfile
            import xml.etree.ElementTree as ET
            import tempfile
            import shutil
            import re

            # 创建专门的子文件夹来存储中间文件
            from pathlib import Path
            input_path = Path(docx_path)

            # 如果文件在Chinese文件夹中，创建processed子文件夹
            if 'Chinese' in str(input_path):
                # 获取Chinese文件夹的路径
                chinese_folder = None
                for parent in input_path.parents:
                    if parent.name == 'Chinese':
                        chinese_folder = parent
                        break

                if chinese_folder:
                    processed_folder = chinese_folder / 'processed'
                    processed_folder.mkdir(exist_ok=True)
                    filename = input_path.name.replace('.docx', '_dot_processed.docx')
                    output_path = str(processed_folder / filename)
                    print(f"中间文件将保存到: processed/{filename}")
                else:
                    # 回退到原来的方式
                    output_path = docx_path.replace('.docx', '_dot_processed.docx')
            else:
                # 不在Chinese文件夹中，使用原来的方式
                output_path = docx_path.replace('.docx', '_dot_processed.docx')

            # 创建临时目录来解压和重新打包docx
            with tempfile.TemporaryDirectory() as temp_dir:
                extract_dir = Path(temp_dir) / "docx_content"
                extract_dir.mkdir()

                # 解压docx文件
                with zipfile.ZipFile(docx_path, 'r') as zip_file:
                    zip_file.extractall(extract_dir)

                # 修改document.xml
                document_xml_path = extract_dir / "word" / "document.xml"
                if document_xml_path.exists():
                    with open(document_xml_path, 'r', encoding='utf-8') as f:
                        xml_content = f.read()

                    # 查找并替换加点字标记
                    run_with_em_pattern = r'(<w:r>.*?<w:rPr>.*?)<w:em w:val="dot"\s*/>(.*?</w:rPr>.*?<w:t>)(.*?)(</w:t>.*?</w:r>)'

                    def replace_run_with_em(match):
                        before_em = match.group(1)
                        after_em = match.group(2)
                        text_content = match.group(3)
                        after_text = match.group(4)

                        # 添加下划线和特殊标记
                        underline_xml = '<w:u w:val="single"/>'
                        marked_text = f"[DOT_BELOW]{text_content}[/DOT_BELOW]"

                        return f"{before_em}{underline_xml}{after_em}{marked_text}{after_text}"

                    # 🌊 安全的波浪线XML预处理（修正版）
                    modified_content = xml_content

                    # 🔧 修复：允许rPr内有其他标签，但确保在同一个w:r内
                    # 关键：只匹配有xml:space="preserve"的wave格式（真正的填空）
                    wavy_pattern = r'(<w:r><w:rPr>(?:[^<]|<[^/][^>]*>)*<w:u w:val="wave"/>(?:[^<]|<[^/][^>]*>)*</w:rPr><w:t[^>]*xml:space="preserve"[^>]*>)(\s+)(</w:t></w:r>)'

                    def replace_wavy_spaces(match):
                        before_text = match.group(1)  # <w:r><w:rPr>...<w:t>
                        spaces = match.group(2)  # 空格内容
                        after_text = match.group(3)  # </w:t></w:r>

                        # 🔍 严格检查：只处理纯空格的波浪线（真正的填空）
                        if spaces.strip() != '':
                            # 如果不是纯空格，则不处理，保持原样
                            return match.group(0)

                        space_count = len(spaces)
                        # 保持XML结构完整，只替换文本内容
                        marked_text = f"[WAVY_SPACE_{space_count}]"
                        return f"{before_text}{marked_text}{after_text}"

                    # 🔍 调试：检查波浪线匹配和处理过程
                    test_matches = re.findall(wavy_pattern, modified_content, flags=re.DOTALL)
                    print(f"  🔍 调试：找到 {len(test_matches)} 个波浪线模式")
                    for i, match in enumerate(test_matches):
                        content = match[1]
                        is_pure_space = content.strip() == ''
                        print(
                            f"    匹配 {i + 1}: 内容长度={len(content)}, 纯空格={is_pure_space}, 内容='{content[:20]}...'")


                # 重新打包docx文件
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for file_path in extract_dir.rglob('*'):
                        if file_path.is_file():
                            relative_path = file_path.relative_to(extract_dir)
                            zip_file.write(file_path, relative_path)

            return output_path

        except Exception as e:
            print(f"  ⚠️ 预处理失败: {e}")
            return None

    def convert_word_to_text(self, file_path, output_format='markdown'):
        """使用pandoc将Word文档转换为文本，并增强格式标注"""
        if not self.pandoc_available:
            print("❌ Pandoc不可用，无法处理文档")
            return None

        if not os.path.exists(file_path):
            print(f"❌ 文件不存在: {file_path}")
            return None

        print(f"开始处理文档: {file_path}")
        print(f"文件大小: {os.path.getsize(file_path) / (1024 * 1024):.2f} MB")

        try:
            # 构建pandoc命令
            import shutil
            from pathlib import Path

            # 获取试卷名（去掉扩展名）
            exam_name = Path(file_path).stem
            exam_media_dir = Path("media") / f"{exam_name}_media"

            # 如果文件夹存在，清空内容
            if exam_media_dir.exists():
                shutil.rmtree(exam_media_dir)
                print(f"🗑️ 清空现有图片文件夹: {exam_media_dir}")
            exam_media_dir.mkdir(parents=True, exist_ok=True)
            print(f"📁 创建图片文件夹: {exam_media_dir}")

            cmd = [
                'pandoc',
                file_path,
                '--to', output_format,
                '--wrap', 'none',  # 不自动换行
                '--standalone',  # 生成独立文档
                '--extract-media', str(exam_media_dir),  # 直接提取到试卷专用文件夹
                '--quiet'  # 减少输出
            ]

            print(f"执行命令: {' '.join(cmd)}")

            # 执行pandoc转换
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)

            if result.returncode == 0:
                content = result.stdout
                print(f"✅ 转换成功: {len(content)} 字符")

                # 如果是docx文件，进行格式分析
                if file_path.lower().endswith('.docx'):
                    print("检测到docx文件，开始格式分析...")
                    # 进行格式分析
                    format_analysis = self.extract_format_analysis(file_path)
                    if format_analysis:
                        print(f"✅ 格式分析完成: {len(self.paragraph_formatting)} 个段落格式")

                # 新增：如果有格式分析结果，增强pandoc内容
                if (hasattr(self, 'special_formatted_text') and self.special_formatted_text) or \
                        (hasattr(self, 'paragraph_formatting') and self.paragraph_formatting):
                    print("🎨 开始增强格式标注...")
                    content = self._enhance_content_with_format_info(content)

                # 处理图片路径：创建试卷名子文件夹并更新图片路径
                content = self._process_image_paths(content, file_path)

                # 保存转换结果
                pandoc_res_dir = Path("pandoc_res")
                pandoc_res_dir.mkdir(exist_ok=True)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = pandoc_res_dir / f"pandoc转换结果_{timestamp}.txt"
                with open(output_filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"转换结果已保存到: {output_filename}")

                return content
            else:
                print(f"❌ 转换失败: {result.stderr}")
                return None

        except subprocess.TimeoutExpired:
            print("❌ 转换超时（5分钟）")
            return None
        except Exception as e:
            print(f"❌ 转换异常: {e}")
            return None

    def _process_image_paths(self, content, file_path):
        """更新图片路径到试卷专用文件夹"""
        import re
        from pathlib import Path

        # 获取试卷名（去掉扩展名）
        exam_name = Path(file_path).stem
        print(f"📁 更新图片路径，试卷名: {exam_name}")

        # 查找所有图片引用并更新路径
        image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        matches = re.findall(image_pattern, content)

        if matches:
            print(f"🖼️ 发现 {len(matches)} 个图片引用")

            for alt_text, image_path in matches:
                # 更新图片路径
                new_path = f"media/{exam_name}_media/{Path(image_path).name}"
                content = content.replace(image_path, new_path)
                print(f"  ✅ 更新路径: {Path(image_path).name}")
        else:
            print("ℹ️ 未发现图片引用")

        return content

    def _clean_dot_below_markers(self, text):
        """清理加点字标记，用于匹配比较"""
        import re

        # 格式1：完整pandoc格式 [\[DOT_BELOW\]字符\[/DOT_BELOW\]]{.underline}
        pattern1 = r'\[\\\[DOT_BELOW\\\]([\u4e00-\u9fff]+)\\\[/DOT_BELOW\\\]\]\{\.underline\}'
        cleaned = re.sub(pattern1, r'\1', text)

        # 格式2：简化格式 [DOT_BELOW]字符[/DOT_BELOW] (支持多个字符)
        pattern2 = r'\[DOT_BELOW\]([\u4e00-\u9fff]+)\[/DOT_BELOW\]'
        cleaned = re.sub(pattern2, r'\1', cleaned)

        # 格式3：处理不完整的DOT_BELOW标记（如截断的文本）
        pattern3 = r'\[DOT_BELOW\]([\u4e00-\u9fff]*)\[/DOT.*?'
        cleaned = re.sub(pattern3, r'\1', cleaned)

        # 格式4：清理剩余的DOT_BELOW开始标记
        pattern4 = r'\[DOT_BELOW\]'
        cleaned = re.sub(pattern4, '', cleaned)

        return cleaned

    def _should_enable_detailed_debug(self, para_text):
        """
        通用调试条件判断：基于文本特征和复杂度决定是否启用详细调试
        """
        # 文本长度相关条件
        text_length = len(para_text)
        if text_length < 5:  # 过短文本通常不需要详细调试
            return False
        if text_length > 50:  # 长文本更需要调试
            return True

        # 包含特殊格式标记
        special_markers = ['DOT_BELOW', '【', '】', '[', ']', '\\[', '\\]']
        if any(marker in para_text for marker in special_markers):
            return True

        # 包含复杂标点或格式
        complex_chars = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩',
                         '“', '”', '‘', '’', '（', '）', '——', '…']
        if any(char in para_text for char in complex_chars):
            return True

        # 包含引号或特殊符号
        if '“' in para_text or "”" in para_text or '「' in para_text or '」' in para_text:
            return True

        # 中等长度的文本，有一定调试价值
        if 20 <= text_length <= 50:
            return True

        return False

    def _has_high_text_similarity(self, text1, text2):
        """
        通用文本相似度判断：基于多种指标计算文本相似度
        """
        if not text1 or not text2:
            return False

        # 长度相似性检查
        len1, len2 = len(text1), len(text2)
        if abs(len1 - len2) > max(len1, len2) * 0.5:  # 长度差异超过50%
            return False

        # 字符集重叠检查
        chars1, chars2 = set(text1), set(text2)
        overlap = len(chars1 & chars2)
        union = len(chars1 | chars2)
        if union > 0:
            char_similarity = overlap / union
            if char_similarity > 0.7:  # 字符重叠度超过70%
                return True

        # 子串包含检查
        shorter, longer = (text1, text2) if len1 < len2 else (text2, text1)
        if len(shorter) >= 5:  # 只对有意义长度的文本做子串检查
            # 检查较短文本的前半部分是否在较长文本中
            half_len = len(shorter) // 2
            if half_len >= 3 and shorter[:half_len] in longer:
                return True
            # 检查较短文本的后半部分是否在较长文本中
            if half_len >= 3 and shorter[-half_len:] in longer:
                return True

        # 中文词汇重叠检查（针对中文文档）
        import re
        chinese_words1 = re.findall(r'[\u4e00-\u9fff]{2,}', text1)
        chinese_words2 = re.findall(r'[\u4e00-\u9fff]{2,}', text2)

        if chinese_words1 and chinese_words2:
            word_overlap = len(set(chinese_words1) & set(chinese_words2))
            word_total = len(set(chinese_words1) | set(chinese_words2))
            if word_total > 0 and word_overlap / word_total > 0.5:  # 词汇重叠超过50%
                return True

        return False

    def _find_centered_text_match(self, para_text, content):
        """专门用于居中文本的匹配方法，优先匹配独立行"""
        lines = content.split('\n')

        # 优先匹配独立行
        for line in lines:
            if line.strip() == para_text:
                return line.strip(), "独立行"

        # 后备：使用原有算法
        return self._find_best_match_in_content(para_text, content)

    def _normalize_quotes(self, text):
        """标准化引号，用于匹配比较"""
        # 将各种中文引号统一为标准引号
        quote_mappings = {
            '“': '"',  # 左双引号 (8220) -> 普通双引号 (34)
            '”': '"',  # 右双引号 (8221) -> 普通双引号 (34)
            '‘': "'",  # 左单引号 (8216) -> 普通单引号 (39)
            '’': "'",  # 右单引号 (8217) -> 普通单引号 (39)
            '「': '"',  # 日式左引号
            '」': '"',  # 日式右引号
            '『': '"',  # 日式左双引号
            '』': '"',  # 日式右双引号
        }

        result = text
        for old_quote, new_quote in quote_mappings.items():
            result = result.replace(old_quote, new_quote)
        return result

    def _find_best_match_in_content(self, para_text, content):
        """在内容中找到段落的最佳匹配位置"""

        # 添加空格处理
        para_text_cleaned = ' '.join(para_text.split())

        # 特殊处理：优先尝试匹配独立行（特别是标题类文本）
        lines = content.split('\n')
        for line in lines:
            line_stripped = line.strip()
            line_cleaned = ' '.join(line_stripped.split())

            # 尝试直接匹配
            if line_stripped == para_text or line_cleaned == para_text_cleaned:
                return line_stripped, "独立行"

            # 尝试标准化后匹配独立行
            normalized_line = self._normalize_quotes(line_stripped)
            normalized_para = self._normalize_quotes(para_text)
            normalized_line_cleaned = ' '.join(normalized_line.split())
            normalized_para_cleaned = ' '.join(normalized_para.split())

            if normalized_line_cleaned == normalized_para_cleaned:
                return line_stripped, "独立行引号清理"

            # 🔧 新增：DOT_BELOW清理后匹配
            cleaned_line = self._clean_dot_below_markers(line_stripped)
            cleaned_para = self._clean_dot_below_markers(para_text)
            if cleaned_line == cleaned_para:
                return line_stripped, "独立行DOT_BELOW清理"

            # 综合处理：DOT_BELOW + 引号 + 空格
            both_cleaned_line = ' '.join(self._normalize_quotes(self._clean_dot_below_markers(line_stripped)).split())
            both_cleaned_para = ' '.join(self._normalize_quotes(self._clean_dot_below_markers(para_text)).split())
            if both_cleaned_line == both_cleaned_para:
                return line_stripped, "独立行综合清理"

            # 🔧 新增：对于短文本（如标题），检查是否作为独立行存在
            if len(para_text) <= 20:  # 标题通常较短
                # 检查该文本是否作为独立行存在（前后都是空行或换行）
                if para_text in line_stripped and len(line_stripped) <= len(para_text) + 5:
                    # 进一步检查：确保不是包含在长句子中
                    if line_stripped == para_text or line_stripped.startswith(para_text) and len(line_stripped) - len(
                            para_text) <= 3:
                        return line_stripped, "独立行短文本"

        # 优化长度策略 - 对短文本更灵活
        if len(para_text) <= 8:
            # 短文本：优先完整匹配，然后逐步减少
            lengths = [len(para_text)]
            if len(para_text) > 3:
                lengths.extend([len(para_text) - 1, len(para_text) - 2])
            if len(para_text) > 5:
                lengths.append(5)
        else:
            # 长文本：使用更多选项
            lengths = [25, 20, 15, 12, 10, 8]

        for length in lengths:
            if len(para_text) < length:
                continue

            para_start = para_text[:length]
            para_start_cleaned = ' '.join(para_start.split())

            # 方法1：直接匹配
            if para_start in content:
                if not any(f"【{marker}】{para_start}" in content for marker in ["首行缩进", "居中", "居右"]):
                    return para_start, f"精确{length}"

            # 方法1.5：空格清理后匹配
            if para_start_cleaned != para_start and para_start_cleaned in content:
                if not any(f"【{marker}】{para_start_cleaned}" in content for marker in ["首行缩进", "居中", "居右"]):
                    return para_start, f"空格清理{length}"

            # 方法2：标准化引号后匹配
            normalized_para_start = self._normalize_quotes(para_start)
            if normalized_para_start != para_start:
                if normalized_para_start in content:
                    if not any(
                            f"【{marker}】{normalized_para_start}" in content for marker in ["首行缩进", "居中", "居右"]):
                        return normalized_para_start, f"引号{length}"

                # 同时标准化引号和清理空格
                normalized_cleaned = ' '.join(normalized_para_start.split())
                if normalized_cleaned != normalized_para_start and normalized_cleaned in content:
                    if not any(f"【{marker}】{normalized_cleaned}" in content for marker in ["首行缩进", "居中", "居右"]):
                        return para_start, f"引号空格{length}"

            # 方法3：清理加点字标记后匹配
            cleaned_para_start = self._clean_dot_below_markers(para_start)
            if cleaned_para_start != para_start:
                if cleaned_para_start in content:
                    if not any(f"【{marker}】{cleaned_para_start}" in content for marker in ["首行缩进", "居中", "居右"]):
                        return cleaned_para_start, f"清理{length}"

                # 🔧 新增：如果直接在content中找不到，尝试逐行匹配
                lines = content.split('\n')
                for line in lines:
                    line_cleaned = self._clean_dot_below_markers(line.strip())
                    if cleaned_para_start in line_cleaned:
                        if not any(f"【{marker}】{cleaned_para_start}" in content for marker in
                                   ["首行缩进", "居中", "居右"]):
                            return cleaned_para_start, f"逐行清理{length}"

            # 方法4：综合处理（引号+加点字+空格）
            both_processed = self._normalize_quotes(self._clean_dot_below_markers(para_start))
            both_processed_cleaned = ' '.join(both_processed.split())

            if both_processed != para_start:
                if both_processed in content:
                    if not any(f"【{marker}】{both_processed}" in content for marker in ["首行缩进", "居中", "居右"]):
                        return both_processed, f"综合{length}"

                if both_processed_cleaned != both_processed and both_processed_cleaned in content:
                    if not any(f"【{marker}】{both_processed_cleaned}" in content for marker in
                               ["首行缩进", "居中", "居右"]):
                        return para_start, f"综合空格{length}"

        # 特殊处理：序号段落（①②③⑭等）
        import re
        if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]', para_text):
            content_without_number = para_text[1:].strip()
            for attempt_length in [len(content_without_number), min(10, len(content_without_number))]:
                if len(content_without_number) >= attempt_length > 0:
                    text_to_find = content_without_number[:attempt_length]
                    if text_to_find in content:
                        if not any(f"【{marker}】{text_to_find}" in content for marker in ["首行缩进", "居中", "居右"]):
                            return para_text, f"序号匹配{attempt_length}"

        # 🆕 回退匹配策略：更宽松的匹配算法
        # 只在调试模式下显示详细信息
        should_debug_fallback = self._should_enable_detailed_debug(para_text)
        if should_debug_fallback:
            print(f"     尝试回退匹配策略...")

        # 回退策略1：模糊字符匹配
        para_chars = set(para_text)
        content_lines = content.split('\n')
        best_match = None
        best_score = 0

        for line in content_lines:
            line_stripped = line.strip()
            if len(line_stripped) < 5:  # 跳过过短的行
                continue

            # 计算字符重叠率
            line_chars = set(line_stripped)
            overlap = len(para_chars & line_chars)
            total_chars = len(para_chars | line_chars)
            if total_chars > 0:
                score = overlap / total_chars

                # 如果重叠率很高且长度相近
                if score > 0.8 and abs(len(line_stripped) - len(para_text)) < max(5, len(para_text) * 0.2):
                    if score > best_score:
                        best_score = score
                        best_match = line_stripped

        if best_match:
            if should_debug_fallback:
                print(f"     ✅ 模糊匹配成功 (相似度: {best_score:.2f})")
            return best_match, f"模糊匹配({best_score:.2f})"

        # 回退策略2：关键词匹配
        # 提取中文字符作为关键词
        chinese_chars = re.findall(r'[\u4e00-\u9fff]+', para_text)
        if chinese_chars:
            # 取最长的中文词汇作为关键词
            key_phrase = max(chinese_chars, key=len)
            if len(key_phrase) >= 3:  # 至少3个字
                for line in content_lines:
                    line_stripped = line.strip()
                    if key_phrase in line_stripped and len(line_stripped) > len(key_phrase):
                        # 检查上下文相似性
                        context_before = para_text[:para_text.find(key_phrase)]
                        context_after = para_text[para_text.find(key_phrase) + len(key_phrase):]

                        line_key_pos = line_stripped.find(key_phrase)
                        line_before = line_stripped[:line_key_pos]
                        line_after = line_stripped[line_key_pos + len(key_phrase):]

                        # 简单的上下文匹配
                        before_match = any(c in line_before for c in context_before[-3:]) if context_before else True
                        after_match = any(c in line_after for c in context_after[:3]) if context_after else True

                        if before_match and after_match:
                            if should_debug_fallback:
                                print(f"     ✅ 关键词匹配成功 (关键词: {key_phrase})")
                            return line_stripped, f"关键词匹配({key_phrase})"

        # 回退策略3：数字序号匹配
        number_match = re.match(r'^([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]|[0-9]+[\.、])', para_text)
        if number_match:
            number_prefix = number_match.group(0)
            remaining_text = para_text[len(number_prefix):].strip()

            for line in content_lines:
                line_stripped = line.strip()
                if remaining_text and len(remaining_text) > 3 and remaining_text[:10] in line_stripped:
                    if should_debug_fallback:
                        print(f"     ✅ 序号内容匹配成功")
                    return line_stripped, "序号内容匹配"

        # 如果所有方法都失败，尝试部分匹配（用于调试）
        escaped_text = re.escape(para_text[:min(5, len(para_text))])
        if re.search(escaped_text, content):
            return None, "部分存在但无法匹配"

        if should_debug_fallback:
            print(f"     ❌ 所有匹配策略都失败")
        return None, None

    def _enhance_content_with_format_info(self, content):
        """根据格式分析结果增强pandoc内容"""
        print(f"🔍 开始分析 {len(self.special_formatted_text)} 个特殊格式文本")
        print(f"📏 开始分析 {len(self.paragraph_formatting)} 个段落格式")

        # 🌊 首先：识别XML预处理标记的波浪线填空（优先级最高）
        import re
        format_enhanced_count = 0

        # 识别XML预处理阶段插入的波浪线空格标记（支持pandoc转义）
        wavy_space_pattern = r'\\?\[WAVY_SPACE_(\d+)\\?\]'
        wavy_space_matches = re.findall(wavy_space_pattern, content)

        if wavy_space_matches:
            print(f"  🌊 发现 {len(wavy_space_matches)} 个XML预处理的波浪线填空标记")

            def replace_wavy_space_marker(match):
                space_count = int(match.group(1))  # 提取真实的空格数量
                nbsp_content = "&nbsp;" * space_count
                enhanced_text = f"[{nbsp_content}]{{.wavy-underline}}"
                print(f"  🌊 转换波浪线填空：原始{space_count}个空格 → {space_count}个&nbsp;")
                return enhanced_text

            # 修正：使用完整模式匹配替换，处理转义的方括号
            content = re.sub(wavy_space_pattern, replace_wavy_space_marker, content)
            format_enhanced_count += len(wavy_space_matches)
        else:
            # 备用：基于格式分析的波浪线填空检测
            wavy_space_texts = []
            for item in self.special_formatted_text:
                if any("波浪线格式" in fmt for fmt in item['formats']):
                    text = item['text']
                    # 检查是否为纯空格文本（可能的填空）
                    if text.replace(' ', '').replace('\u00A0', '') == '':
                        wavy_space_texts.append((text, len(text)))
                        print(f"  🌊 格式分析发现波浪线填空：{len(text)}个空格")

            # 处理检测到的波浪线填空
            for original_text, space_count in wavy_space_texts:
                nbsp_content = "&nbsp;" * space_count
                enhanced_text = f"[{nbsp_content}]{{.wavy-underline}}"
                print(f"  🌊 转换波浪线填空：{space_count}个&nbsp;")
                content = content.replace(original_text, enhanced_text, 1)
                format_enhanced_count += 1

            if not wavy_space_texts:
                print(f"  ⚠️ 未检测到波浪线填空，跳过处理")

        # ⚠️ 注意：删除重复的波浪线处理逻辑，使用通用的格式标注处理机制

        # 🚨 重要：优先处理段落格式（居右、居中、首行缩进），避免被特殊格式文本处理干扰
        # 第一步：处理居右段落
        right_aligned_enhanced_count = 0
        for para_info in self.paragraph_formatting:
            if para_info['is_right_aligned']:
                para_text = para_info['text'].strip()
                print(f"🔍 检查居右文本: \"{para_text}\" (长度: {len(para_text)})")

                # 跳过过短的文本
                if len(para_text) < 1:
                    continue

                # 使用改进的匹配算法
                match_result, match_type = self._find_best_match_in_content(para_text, content)

                if match_result:
                    # 在匹配的文本前添加标记
                    enhanced_start = f"【居右】{match_result}"
                    content = content.replace(match_result, enhanced_start, 1)
                    right_aligned_enhanced_count += 1
                    print(f"✅ 居右标记({match_type}): \"{match_result[:30]}...\"")
                else:
                    print(f"❌ 居右未匹配: \"{para_text[:30]}...\" (长度: {len(para_text)})")

        # 第二步：处理居中段落
        centered_enhanced_count = 0
        for para_info in self.paragraph_formatting:
            if para_info['is_centered']:
                para_text = para_info['text'].strip()

                # 跳过过短的文本
                if len(para_text) < 2:
                    continue

                # 避免重复标记（如果已经有居右标记）
                check_lengths = [min(10, len(para_text)), min(8, len(para_text))] if len(para_text) > 5 else [
                    len(para_text)]
                if any(f"【居右】{para_text[:length]}" in content for length in check_lengths if
                       len(para_text) >= length):
                    print(f"  → 跳过：已有居右标记")
                    continue

                # 🔧 优化：对于居中文本，优先寻找独立行匹配
                match_result, match_type = self._find_centered_text_match(para_text, content)

                if match_result:
                    # 对于独立行，使用行级别替换
                    if match_type == "独立行":
                        lines = content.split('\n')
                        for i, line in enumerate(lines):
                            if line.strip() == match_result:
                                lines[i] = f"【居中】{line}"
                                content = '\n'.join(lines)
                                break
                    else:
                        # 其他情况使用普通替换
                        enhanced_start = f"【居中】{match_result}"
                        content = content.replace(match_result, enhanced_start, 1)

                    centered_enhanced_count += 1
                    print(f"✅ 居中标记({match_type}): \"{match_result[:30]}...\"")
                else:
                    print(f"❌ 居中未匹配: \"{para_text[:20]}...\"")

        # 第三步：处理段落首行缩进（最后处理，避免误抢其他格式）
        indent_enhanced_count = 0
        for para_info in self.paragraph_formatting:
            if para_info['has_first_line_indent']:
                para_text = para_info['text'].strip()

                # 跳过过短的文本
                if len(para_text) < 8:
                    continue

                # 避免重复标记（如果已经有其他标记）
                check_lengths = [min(10, len(para_text)), min(8, len(para_text))] if len(para_text) > 5 else [
                    len(para_text)]
                if any(f"【居右】{para_text[:length]}" in content for length in check_lengths if
                       len(para_text) >= length):
                    print(f"  → 跳过：已有居右标记")
                    continue
                if any(f"【居中】{para_text[:length]}" in content for length in check_lengths if
                       len(para_text) >= length):
                    print(f"  → 跳过：已有居中标记")
                    continue

                # 使用改进的匹配算法
                match_result, match_type = self._find_best_match_in_content(para_text, content)

                if match_result:
                    # 在匹配的文本前添加标记
                    enhanced_start = f"【首行缩进】{match_result}"
                    content = content.replace(match_result, enhanced_start, 1)
                    indent_enhanced_count += 1
                    print(f"✅ 缩进标记({match_type}): \"{match_result[:30]}...\"")
                else:
                    print(f"❌ 缩进未匹配: \"{para_text[:30]}...\"")
                    # 通用调试条件：基于文本特征和复杂度判断是否需要详细调试
                    should_debug = self._should_enable_detailed_debug(para_text)
                    if should_debug:
                        print(f"  → 详细调试匹配过程:")
                        print(f"     原文本: {repr(para_text[:80])}")
                        print(f"     文本长度: {len(para_text)}")

                        # 显示各种清理步骤
                        cleaned_text = self._clean_dot_below_markers(para_text)
                        normalized_text = self._normalize_quotes(para_text)
                        fully_cleaned = self._normalize_quotes(cleaned_text)

                        print(f"     DOT_BELOW清理后: {repr(cleaned_text[:80])}")
                        print(f"     引号标准化后: {repr(normalized_text[:80])}")
                        print(f"     完全清理后: {repr(fully_cleaned[:80])}")

                        # 检查各种匹配可能性
                        content_lines = content.split('\n')
                        found_similar = []

                        for i, line in enumerate(content_lines):
                            line_stripped = line.strip()
                            if not line_stripped:
                                continue

                            # 检查各种匹配（移除硬编码的内容判断）
                            matches = []
                            if cleaned_text in line_stripped:
                                matches.append("DOT_BELOW清理")
                            if fully_cleaned in line_stripped:
                                matches.append("完全清理")
                            if para_text[:20] in line_stripped:
                                matches.append("前20字符")
                            if line_stripped[:20] in para_text:
                                matches.append("行前20字符")

                            # 通用相似度检查，移除具体内容判断
                            if matches or self._has_high_text_similarity(para_text, line_stripped):
                                found_similar.append(
                                    f"     第{i + 1}行: {repr(line_stripped[:80])} [{', '.join(matches) if matches else '高相似度'}]")

                        if found_similar:
                            print(f"     找到相似内容:")
                            for similar in found_similar[:3]:  # 只显示前3个
                                print(similar)
                        else:
                            print(f"     ❌ 未找到任何相似内容")

                        # 尝试模糊匹配
                        for j, line in enumerate(content_lines):
                            line_stripped = line.strip()
                            if len(line_stripped) > 10 and abs(len(line_stripped) - len(para_text)) < 10:
                                # 计算相似度（简单的字符重叠）
                                overlap = len(set(para_text) & set(line_stripped))
                                if overlap > len(para_text) * 0.7:
                                    print(
                                        f"     🔍 高相似度行: {repr(line_stripped[:60])} (重叠字符: {overlap}/{len(para_text)})")
                                    cleaned_line = self._clean_dot_below_markers(line)
                                    print(f"     第{j + 1}行清理后: {repr(cleaned_line[:80])}")

        # 按文本长度排序，从长到短，避免短文本替换影响长文本
        sorted_formats = sorted(self.special_formatted_text,
                                key=lambda x: len(x['text']), reverse=True)

        format_enhanced_count = 0

        for format_item in sorted_formats:
            text = format_item['text'].strip()
            formats = format_item['formats']

            # 跳过空文本或过短文本
            if len(text) < 2:
                continue

            # 跳过包含居中文本的特殊格式
            for para_info in self.paragraph_formatting:
                if para_info['is_centered'] and para_info['text'] in text and para_info['text'] != text:
                    print(f"  ⏭️ 跳过包含居中文本的特殊格式: \"{text[:30]}...\"")
                    continue

            # 生成格式标注
            format_annotation = self._generate_format_annotation(formats)

            # 🔧 字符标准化：处理pandoc转换导致的字符差异
            # 中文引号 → 英文引号，中文省略号 → 英文点号
            normalized_text = text.replace('“', '"').replace('”', '"').replace('……', '......')

            # 🛡️ 防重复处理：检查文本是否已经被标记过
            if format_annotation:

                # 智能匹配：先尝试原始文本，再尝试标准化文本
                target_text = None
                if text in content:
                    target_text = text
                elif normalized_text in content:
                    target_text = normalized_text

                if target_text:
                    # 检查是否已被居中标记
                    if f"【居中】{target_text}" in content:
                        print(f"  ⏭️ 跳过已被居中标记的文本: \"{target_text[:30]}...\"")
                        continue

                    # 创建增强的文本标注
                    enhanced_text = f"[{target_text}]{{{format_annotation}}}"
                    content = content.replace(target_text, enhanced_text, 1)
                    format_enhanced_count += 1

                    if target_text != text:
                        print(
                            f"🔧 字符标准化成功: \"{text[:30]}{'...' if len(text) > 30 else ''}\" -> {format_annotation}")
                    else:
                        print(f"格式增强: \"{text[:30]}{'...' if len(text) > 30 else ''}\" -> {format_annotation}")

        print(f"✅ 格式增强完成:")
        print(f" 居右对齐标记: {right_aligned_enhanced_count} 个段落")
        print(f" 居中对齐标记: {centered_enhanced_count} 个段落")
        print(f" 首行缩进标记: {indent_enhanced_count} 个段落")
        print(f" 特殊格式标记: {format_enhanced_count} 个文本")
        return content

    def _generate_format_annotation(self, formats):
        """根据格式列表生成标注"""
        annotations = []

        # 🌊 特殊处理：波浪线格式优先，只返回波浪线标记
        for fmt in formats:
            if "波浪线格式" in fmt:
                return ".wavy-underline"  # 直接返回，不添加其他格式

        # 处理其他下划线类型
        for fmt in formats:
            if "点状线格式" in fmt:
                annotations.append(".dotted-underline")
            elif "下划线: 单下划线" in fmt:
                annotations.append(".single-underline")
            elif "下划线: 双下划线" in fmt:
                annotations.append(".double-underline")
            elif "下划线: 粗下划线" in fmt:
                annotations.append(".thick-underline")
            elif "下划线: 虚线下划线" in fmt:
                annotations.append(".dashed-underline")
            elif "删除线" in fmt:
                annotations.append(".strikethrough")
            elif "上标" in fmt:
                annotations.append(".superscript")
            elif "下标" in fmt:
                annotations.append(".subscript")
            elif "着重号" in fmt:
                annotations.append(".emphasis-mark")
            elif "粗体" in fmt:
                annotations.append(".bold")
            elif "斜体" in fmt:
                annotations.append(".italic")

        # 处理字体颜色（提取颜色值）
        for fmt in formats:
            if "字体颜色:" in fmt and "000000" not in fmt:  # 跳过黑色
                color = fmt.split("字体颜色:")[-1].strip()
                annotations.append(f".color-{color}")

        # 处理字号
        for fmt in formats:
            if "字号:" in fmt and "磅" in fmt:
                size = fmt.split("字号:")[-1].replace("磅", "").strip()
                try:
                    size_num = float(size)
                    if size_num != 12.0:  # 跳过默认字号
                        annotations.append(f".font-{size}pt")
                except:
                    pass

        return " ".join(annotations) if annotations else None

    def _convert_dashes_to_chinese(self, content):
        """转换连续短横线为中文破折号"""
        print("🔀 转换连续短横线为中文破折号...")

        try:
            import re

            conversion_count = 0

            dash_pattern = r'-{3,}'  # 匹配3个或更多连续的短横线

            def replace_dashes(match):
                nonlocal conversion_count
                dashes = match.group(0)
                dash_count = len(dashes)
                conversion_count += 1
                # 每3个短横线替换为一个em dash
                em_dash_count = dash_count // 3
                return '—' * em_dash_count

            content = re.sub(dash_pattern, replace_dashes, content)

            if conversion_count > 0:
                print(f"  ✅ 转换了 {conversion_count} 处连续短横线为中文破折号（每3个短横线转换为1个em dash）")
            else:
                print("  ℹ️ 未发现需要转换的连续短横线")

            return content

        except Exception as e:
            print(f"  ⚠️ 破折号转换失败: {e}")
            return content

    def call_llm_api(self, content, prompt_template_path="prompt.md"):
        """调用大模型API解析文档结构"""
        print("开始调用大模型API...")

        # 读取prompt模板
        try:
            with open(prompt_template_path, 'r', encoding='utf-8') as f:
                prompt_template = f.read()
            # 使用安全的字符串替换
            prompt = prompt_template.replace("{content}", content)
            print(f"成功加载prompt模板: {prompt_template_path}")
        except FileNotFoundError:
            print(f"❌ 未找到prompt模板文件: {prompt_template_path}")
            print("❌ 这是一个严重错误！必须使用正确的prompt模板！")
            print("❌ 默认prompt与优化后的要求不匹配，会导致选项缺失等问题")
            print("💡 请确保prompt_Chinese.md文件存在且可读")
            return None

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }

        # 先测试API连接
        test_data = {
            "model": "doubao-seed-1-6-250615",
            "messages": [{"role": "user", "content": "Hello"}],
            "max_tokens": 32000,
            "temperature": 0.1,
            "stream": True,
            "thinking": {
                "type": "enabled",
                "budget_tokens": 2000
            }
        }

        print("🔍 测试API连接...")
        try:
            test_response = requests.post(
                f"{self.base_url}/v3/chat/completions",
                headers=headers,
                json=test_data,
                timeout=30
            )
            test_response.raise_for_status()
            print("✅ API连接测试成功")
        except requests.exceptions.HTTPError as e:
            print(f"❌ API连接测试失败: {e}")
            print(f"响应状态码: {test_response.status_code}")
            print(f"响应内容: {test_response.text}")
            return None
        except Exception as e:
            print(f"❌ API连接测试异常: {e}")
            return None

        # 正式请求 - 使用流式输出
        data = {
            "model": "doubao-seed-1-6-250615",
            "messages": [{"role": "user", "content": prompt}],
            "stream": True,
            "thinking": {
                "type": "enabled",
                "budget_tokens": 1500
            },
            "response_format": {
                "type": "json_object"
            },
            "temperature": 0.1,
            "max_completion_tokens": 32000
        }

        print("🚀 调用大模型API解析文档结构...")
        print(f"请求URL: {self.base_url}/v3/chat/completions")
        print(f"模型: {data['model']}")
        print(f"消息长度: {len(prompt)} 字符")

        try:
            response = requests.post(
                f"{self.base_url}/v3/chat/completions",
                headers=headers,
                json=data,
                timeout=300,  # 减少超时时间到5分钟
                stream=True
            )
            response.raise_for_status()

            # 处理流式响应
            llm_content = ""
            print("📡 开始接收流式响应...")

            for line in response.iter_lines():
                if line:
                    line = line.decode('utf-8')
                    if line.startswith('data: '):
                        data_str = line[6:]  # 去掉 'data: ' 前缀

                        if data_str == '[DONE]':
                            print("\n✅ 流式响应接收完成")
                            break

                        try:
                            data_json = json.loads(data_str)
                            if 'choices' in data_json and len(data_json['choices']) > 0:
                                choice = data_json['choices'][0]

                                # 处理thinking状态
                                if 'thinking' in choice:
                                    thinking = choice['thinking']
                                    if thinking.get('type') == 'thinking':
                                        print(f"🤔 思考中... ({thinking.get('tokens_used', 0)} tokens)")
                                    elif thinking.get('type') == 'finished':
                                        print(f"✅ 思考完成，共使用 {thinking.get('tokens_used', 0)} tokens")

                                # 处理delta内容
                                if 'delta' in choice and 'content' in choice['delta']:
                                    content = choice['delta']['content']
                                    llm_content += content
                                    print(content, end='', flush=True)

                        except json.JSONDecodeError:
                            continue

            print(f"\n✅ API调用成功，响应长度: {len(llm_content)} 字符")
            return llm_content

        except requests.exceptions.HTTPError as e:
            print(f"HTTP错误: {e}")
            print(f"响应状态码: {response.status_code}")
            print(f"响应内容: {response.text}")
            return None
        except Exception as e:
            print(f"❌ API调用异常: {e}")
            return None


    def _process_api_response(self, llm_content, original_file_path):
        """处理API响应并保存结果，增强健壮性和错误处理"""

        def extract_json_from_codeblock(content):
            """从Markdown代码块中提取JSON内容，处理多种格式情况"""
            # 匹配```json开头的代码块（允许前后有空白）
            json_block_pattern = re.compile(r'^\s*```\s*json\s*\n(.*?)\n\s*```\s*$', re.DOTALL | re.MULTILINE)
            match = json_block_pattern.search(content)
            if match:
                return match.group(1).strip()

            # 匹配普通```代码块
            general_block_pattern = re.compile(r'^\s*```\s*\n(.*?)\n\s*```\s*$', re.DOTALL | re.MULTILINE)
            match = general_block_pattern.search(content)
            if match:
                return match.group(1).strip()

            # 无代码块时返回原始内容（可能已是纯JSON）
            return content.strip()

        def clean_json_string(content):
            """清理JSON字符串，处理常见格式问题"""
            # 移除JSON中的注释（/* ... */ 或 // ...）
            content = re.sub(r'/\*.*?\*/', '', content, flags=re.DOTALL)
            content = re.sub(r'//.*?$', '', content, flags=re.MULTILINE)

            # 移除尾逗号（如 [1,2,] 或 {"a":1,}）
            content = re.sub(r',\s*([}\]])', r'\1', content)

            return content.strip()

        # 1. 提取并清理内容
        try:
            # 提取JSON内容（处理代码块）
            extracted_content = extract_json_from_codeblock(llm_content)

            # 清理JSON格式问题
            cleaned_content = clean_json_string(extracted_content)
        except Exception as e:
            print(f"⚠️ 内容提取/清理失败: {e}")
            cleaned_content = llm_content.strip()

        # 2. 保存原始响应（无论后续处理是否成功，便于调试）
        try:
            # 创建专门的原始响应目录
            raw_dir = Path("raw_api_responses")
            raw_dir.mkdir(exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            raw_filename = raw_dir / f"raw_response_{Path(original_file_path).stem}_{timestamp}.txt"
            with open(raw_filename, 'w', encoding='utf-8') as f:
                f.write(f"=== 原始LLM响应 ===\n{llm_content}\n\n")
                f.write(f"=== 提取后内容 ===\n{extracted_content}\n\n")
                f.write(f"=== 清理后内容 ===\n{cleaned_content}")
            print(f"📄 原始响应已保存到: {raw_filename}")
        except Exception as e:
            print(f"⚠️ 保存原始响应失败: {e}")

        # 3. 解析JSON并保存结果
        try:
            questions = cleaned_content

            # 验证解析结果格式
            if not isinstance(questions, list):
                raise ValueError("解析结果不是JSON数组")

            # 保存处理后的结果
            json_res_dir = Path("json_res")
            json_res_dir.mkdir(exist_ok=True)

            output_file = json_res_dir / f"questions_{Path(original_file_path).stem}_{timestamp}.json"
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(questions, f, ensure_ascii=False, indent=2)

            print(f"🎉 完成！共{len(questions)}道题目，保存到: {output_file}")
            return questions

        except json.JSONDecodeError as e:
            print(f"❌ JSON解析失败: {e}")
            print(f"错误位置: 第{e.lineno}行第{e.colno}列")
            return None
        except ValueError as e:
            print(f"❌ 解析结果格式错误: {e}")
            return None
        except Exception as e:
            print(f"❌ 处理响应时发生意外错误: {e}")
            return None

    def call_coze_workflow(self, processed_data):
        """调用Coze工作流"""
        print("🔗 准备调用Coze工作流...")

        try:
            headers = {
                'Authorization': f'Bearer pat_Z0r3WQNZ435IUDhJCc0bVHDd9mVcIh0Z6tOvYd3HPT3Q6WNfw5KaX7veOhNkqC3N',
                'Content-Type': 'application/json'
            }

            data = {
                "workflow_id": "7540878860784680995",
                "parameters": {
                    "input": json.dumps(processed_data, ensure_ascii=False)
                }
            }

            print("🚀 调用Coze工作流开始...")
            print(f"📊 发送数据量: {len(json.dumps(processed_data, ensure_ascii=False))} 字符")

            response = requests.post('https://api.coze.cn/v1/workflow/run', headers=headers, data=json.dumps(data))

            if response.status_code == 200:
                response_data = response.json().get("data")

                if response_data:
                    # 解析 JSON 字符串
                    parsed_data = json.loads(response_data)

                    # 提取 data 字段并按 \n 分割成数组
                    id_list = parsed_data["data"].split("\n")

                    print(f"✅ Coze工作流调用成功，返回 {len(id_list)} 个ID")
                    print(f"📋 ID列表预览: {', '.join(id_list[:5])}...")  # 只显示前5个

                    return id_list
                else:
                    print("❌ Coze工作流返回数据为空")
                    return None
            else:
                print(f"❌ Coze工作流调用失败，状态码: {response.status_code}")
                print(f"错误信息: {response.text}")
                return None

        except Exception as e:
            print(f"❌ 调用Coze工作流异常: {e}")
            return None

    # 主流程函数
    def process_word_document(self, file_path, output_format='markdown', prompt_template_path="prompt.md",
                              enable_dot_below_detection=True, enable_coze_workflow=True):

        """提示性信息"""
        print("=" * 60)
        print("Pandoc Word文档处理工具 - 增强版 (支持加点字)")
        print("=" * 60)
        print(f"文档文件: {file_path}")
        print(f"输出格式: {output_format}")
        print(f"Prompt模板: {prompt_template_path}")
        # print(f"格式分析: 已禁用")  # 已移除格式分析功能
        print(f"加点字标记保留: {'启用' if enable_dot_below_detection else '禁用'}")
        print(f"Coze工作流: {'启用' if enable_coze_workflow else '禁用'}")
        print("=" * 60)
        """提示结束"""

        # 第一步 - 加点字和波浪线预处理（如果启用且为docx文件）
        processed_file_path = file_path
        if enable_dot_below_detection and file_path.lower().endswith('.docx'):
            processed_file_path = self.preprocess(file_path)
            if not processed_file_path:
                processed_file_path = file_path  # 回退到原文件

        """打印文本中一些格式信息，例如：缩进信息、对齐方式（debug类型的函数，与功能无关）"""
        # 第二步 - 格式分析（如果启用且为docx文件）
        format_analysis = None
        if processed_file_path.lower().endswith('.docx'):
            format_analysis = self.extract_format_analysis(processed_file_path)

        # 第三步：使用pandoc转换文档
        content = self.convert_word_to_text(processed_file_path, output_format)
        if not content:
            print("❌ 文档转换失败，无法继续处理")
            return None

        # 第四步：转换连续短横线为中文破折号
        content = self._convert_dashes_to_chinese(content)

        # 第五步：调用大模型API解析内容
        llm_response = self.call_llm_api(content, prompt_template_path)
        if not llm_response:
            print("❌ API调用失败")
            return None

        # 第六步：处理API响应并集成格式信息
        api_result = self._process_api_response(llm_response, file_path)

        # 第七步：调用Coze工作流（如果启用）
        coze_ids = None
        if enable_coze_workflow:
            print("\n" + "=" * 60)
            print("🔗 Coze工作流处理阶段")
            print("=" * 60)

            if api_result:
                # 正常情况：使用API解析结果调用Coze
                coze_ids = self.call_coze_workflow(api_result)

                if coze_ids:
                    # 创建coze_res文件夹
                    coze_res_dir = Path("coze_res")
                    coze_res_dir.mkdir(exist_ok=True)

                    # 将Coze返回的ID列表保存为文本文件
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    coze_output_file = coze_res_dir / f"coze_ids_{Path(file_path).stem}_{timestamp}.txt"

                    with open(coze_output_file, 'w', encoding='utf-8') as f:
                        f.write(",".join(coze_ids))

                    print(f"📁 Coze ID列表已保存到: {coze_output_file}")

                    # 可选：将ID列表添加到API结果中
                    if isinstance(api_result, list):
                        # 如果API结果是题目列表，可以为每道题添加一个ID
                        for i, question in enumerate(api_result[:len(coze_ids)]):
                            if isinstance(question, dict):
                                question['coze_id'] = coze_ids[i] if i < len(coze_ids) else None
                    print("✅ Coze ID已整合到题目结果中")
                else:
                    print("⚠️ Coze工作流未返回有效数据")
            else:
                # API解析失败的情况：提供手动调用指导
                print("⚠️ 由于JSON解析失败，无法自动调用Coze工作流")

        return {
            'questions': api_result,
            'coze_ids': coze_ids
        } if enable_coze_workflow else api_result


def main():
    """主函数"""
    import sys

    # 配置参数
    if len(sys.argv) > 1:
        word_file_path = sys.argv[1]
    else:
        word_file_path = "精品解析：2025年山东省泰安市中考英语真题 （解析版）.docx"  # 默认文件路径

    output_format = "markdown"  # 可选: markdown, plain, html
    prompt_template_path = "prompt_English.md"

    # 创建处理器实例
    processor = PandocWordProcessor()

    # 检查pandoc可用性
    if not processor.pandoc_available:
        print("❌ Pandoc不可用，请先安装pandoc")
        print("安装方法:")
        print("  macOS: brew install pandoc")
        print("  Ubuntu/Debian: sudo apt-get install pandoc")
        print("  Windows: 下载安装包 https://pandoc.org/installing.html")
        return

    # 处理文档
    result = processor.process_word_document(
        word_file_path,
        output_format,
        prompt_template_path
    )

    if result:
        print("✅ 文档处理完成！")
    else:
        print("❌ 文档处理失败")


if __name__ == "__main__":
    main()

"""

整体思路：三大模块：

    1. 使用pandoc库以及一些工具函数对传入的word进行各种格式解析和处理，最终返回一个markdown格式的文本（string）

    2. 调用大模型API，传入markdown格式的文本，返回大模型对markdown文本处理后的原始相应内容（string），调用_process_api_response函数

    对原始响应内容做后处理，处理成最终正确格式的json内容

    3. 把第二步的json信息，传入coze工作流进行处理，得到最终的题目id


"""